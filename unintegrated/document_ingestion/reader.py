import json
import logging
from abc import ABC
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Optional, List, Dict, Any, ClassVar, Type

import docx
import pymupdf  # PyMuPDF for PDF handling
import pandas as pd
from PIL import Image
from bs4 import BeautifulSoup
import xml.etree.ElementTree as ET
from io import BytesIO


from file import File, FileObject, FileExtensions

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class ReadConfig:
  """Configuration for file reading operations."""
  extract_tables: bool = True
  extract_images: bool = True
  max_workers: int = 4
  image_size_limit: int = 10 * 1024 * 1024  # 10MB
  table_row_limit: int = 10000
  ocr_enabled: bool = False  # For future implementation
  chunk_size: int = 1024 * 1024  # 1MB for chunked reading

class FileReadError(Exception):
  """Custom exception for file reading errors."""
  pass

class FileReader:
  # Class variables for supported formats and their corresponding MIME types
  MIME_TYPES: ClassVar[Dict[str, str]] = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "doc": "application/msword",
    "txt": "text/plain",
    "csv": "text/csv",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "xls": "application/vnd.ms-excel",
    "jpg": "image/jpeg",
    "png": "image/png",
    "svg": "image/svg+xml",
    "html": "text/html",
    "xml": "application/xml",
    "json": "application/json"
  }

  @classmethod
  def read_file(cls, file_path: str, config: Optional[ReadConfig] = None) -> FileObject:
    """
    Read a file and return a FileObject containing its contents.
    
    Args:
      file_path: Path to the file to read
      config: Optional configuration for reading
    
    Returns:
      FileObject containing the file's contents
    
    Raises:
      FileReadError: If there's an error reading the file
      ValueError: If the file type is not supported
    """
    config = config or ReadConfig()
    path = Path(file_path)
    
    if not path.exists():
      raise FileReadError(f"File not found: {file_path}")
    
    # Get file extension and validate
    extension = path.suffix.lower()[1:]
    if extension not in FileExtensions.__members__.values():
      raise ValueError(f"Unsupported file type: {extension}")
    
    try:
      # Get the appropriate reader method
      reader_method = cls._get_reader_method(extension)
      return reader_method(path, config)
    except Exception as e:
      logger.error(f"Error reading file {file_path}: {str(e)}")
      raise FileReadError(f"Error reading file: {str(e)}")

  @classmethod
  def _get_reader_method(cls, extension: str) -> callable:
    """Get the appropriate reader method for the file extension."""
    readers = {
      FileExtensions.PDF.value: cls._read_pdf,
      FileExtensions.DOCX.value: cls._read_docx,
      FileExtensions.DOC.value: cls._read_doc,
      FileExtensions.TXT.value: cls._read_txt,
      FileExtensions.CSV.value: cls._read_csv,
      FileExtensions.XLSX.value: cls._read_xlsx,
      FileExtensions.XLS.value: cls._read_xls,
      FileExtensions.JPG.value: cls._read_image,
      FileExtensions.PNG.value: cls._read_image,
      FileExtensions.SVG.value: cls._read_svg,
      FileExtensions.HTML.value: cls._read_html,
      FileExtensions.XML.value: cls._read_xml,
      FileExtensions.JSON.value: cls._read_json
    }
    return readers[extension]

  @staticmethod
  def _read_pdf(path: Path, config: ReadConfig) -> FileObject:
    """Read PDF files."""
    file_obj = FileObject()
    doc: pymupdf.Document = pymupdf.open(path)
    
    with ThreadPoolExecutor(max_workers=config.max_workers) as executor:
      # Process pages in parallel
      for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Extract text
        text = page.get_text()
        if text.strip():
          file_obj.text_sections.append(text)
        
        if config.extract_tables:
          # Extract tables using PDF table extraction
          tables = page.get_tables()
          if tables:
            file_obj.tables.extend([
              pd.DataFrame(table).to_dict()
              for table in tables
              if len(table) <= config.table_row_limit
            ])
        
        if config.extract_images:
          # Extract images
          image_list = page.get_images(full=True)
          for img_index, img_info in enumerate(image_list):
            try:
              xref = img_info[0]
              base_image = doc.extract_image(xref)
              if base_image and base_image["size"] <= config.image_size_limit:
                file_obj.images.append(base_image["image"])
            except Exception as e:
              logger.warning(f"Error extracting image {img_index} from page {page_num}: {str(e)}")
    
    return file_obj

  @staticmethod
  def _read_docx(path: Path, config: ReadConfig) -> FileObject:
    """Read DOCX files."""
    file_obj = FileObject()
    doc = docx.Document(path)
    
    # Extract text
    for paragraph in doc.paragraphs:
      if paragraph.text.strip():
        file_obj.text_sections.append(paragraph.text)
    
    if config.extract_tables:
      # Extract tables
      for table in doc.tables:
        data = []
        for row in table.rows:
          row_data = [cell.text for cell in row.cells]
          data.append(row_data)
        if data and len(data) <= config.table_row_limit:
          file_obj.tables.append(pd.DataFrame(data).to_dict())
    
    return file_obj

  @staticmethod
  def _read_txt(path: Path, config: ReadConfig) -> FileObject:
    """Read text files."""
    file_obj = FileObject()
    
    with open(path, 'r', encoding='utf-8') as f:
      while True:
        chunk = f.read(config.chunk_size)
        if not chunk:
          break
        if chunk.strip():
          file_obj.text_sections.append(chunk)
    
    return file_obj

  @staticmethod
  def _read_csv(path: Path, config: ReadConfig) -> FileObject:
    """Read CSV files."""
    file_obj = FileObject()
    
    try:
      df = pd.read_csv(path, nrows=config.table_row_limit)
      file_obj.tables.append(df.to_dict())
      file_obj.text_sections.append(df.to_string())
    except pd.errors.EmptyDataError:
      logger.warning(f"Empty CSV file: {path}")
    
    return file_obj

  @staticmethod
  def _read_xlsx(path: Path, config: ReadConfig) -> FileObject:
    """Read XLSX files."""
    file_obj = FileObject()
    
    xlsx = pd.ExcelFile(path)
    for sheet_name in xlsx.sheet_names:
      df = pd.read_excel(xlsx, sheet_name=sheet_name, nrows=config.table_row_limit)
      if not df.empty:
        file_obj.tables.append({
          'sheet_name': sheet_name,
          'data': df.to_dict()
        })
        file_obj.text_sections.append(f"Sheet: {sheet_name}\n{df.to_string()}")
    
    return file_obj

  @staticmethod
  def _read_image(path: Path, config: ReadConfig) -> FileObject:
    """Read image files."""
    file_obj = FileObject()
    
    if path.stat().st_size <= config.image_size_limit:
      with Image.open(path) as img:
        # Convert image to bytes
        img_byte_arr = BytesIO()
        img.save(img_byte_arr, format=img.format)
        file_obj.images.append(img_byte_arr.getvalue())
    
    return file_obj

  @staticmethod
  def _read_html(path: Path, config: ReadConfig) -> FileObject:
    """Read HTML files."""
    file_obj = FileObject()
    
    with open(path, 'r', encoding='utf-8') as f:
      soup = BeautifulSoup(f.read(), 'html.parser')
      
      # Extract text
      text = soup.get_text(separator='\n', strip=True)
      if text:
        file_obj.text_sections.append(text)
      
      if config.extract_tables:
        # Extract tables
        tables = soup.find_all('table')
        for table in tables:
          rows = []
          for row in table.find_all('tr'):
            cols = row.find_all(['td', 'th'])
            if cols:
              rows.append([col.get_text(strip=True) for col in cols])
          if rows and len(rows) <= config.table_row_limit:
            file_obj.tables.append(pd.DataFrame(rows).to_dict())
    
    return file_obj

  @staticmethod
  def _read_xml(path: Path, config: ReadConfig) -> FileObject:
    """Read XML files."""
    file_obj = FileObject()
    
    tree = ET.parse(path)
    root = tree.getroot()
    
    def extract_text(element: ET.Element) -> str:
      """Recursively extract text from XML elements."""
      text = element.text or ''
      for child in element:
        text += extract_text(child)
      text += element.tail or ''
      return text
    
    # Extract text content
    text = extract_text(root)
    if text.strip():
      file_obj.text_sections.append(text)
    
    return file_obj

  @staticmethod
  def _read_json(path: Path, config: ReadConfig) -> FileObject:
    """Read JSON files."""
    file_obj = FileObject()
    
    with open(path, 'r', encoding='utf-8') as f:
      data = json.load(f)
      
      # Convert to string representation
      text = json.dumps(data, indent=2)
      if text.strip():
        file_obj.text_sections.append(text)
      
      # If the JSON contains arrays of objects, treat them as tables
      if isinstance(data, list) and data and isinstance(data[0], dict):
        if len(data) <= config.table_row_limit:
          file_obj.tables.append(pd.DataFrame(data).to_dict())
    
    return file_obj

  @staticmethod
  def _read_doc(path: Path, config: ReadConfig) -> FileObject:
    """
    Read DOC files.
    Note: This is a placeholder. You might want to use a third-party library
    like antiword or python-docx-reader for actual implementation.
    """
    raise NotImplementedError("DOC file reading is not implemented")

  @staticmethod
  def _read_xls(path: Path, config: ReadConfig) -> FileObject:
    """Read XLS files."""
    return FileReader._read_xlsx(path, config)  # Use the same method as xlsx

  @staticmethod
  def _read_svg(path: Path, config: ReadConfig) -> FileObject:
    """Read SVG files."""
    file_obj = FileObject()
    
    if path.stat().st_size <= config.image_size_limit:
      with open(path, 'rb') as f:
        file_obj.images.append(f.read())
    
    return file_obj

  @classmethod
  def get_mime_type(cls, extension: str) -> str:
    """Get MIME type for a file extension."""
    return cls.MIME_TYPES.get(extension.lower(), "application/octet-stream")

  @classmethod
  def is_supported_extension(cls, extension: str) -> bool:
    """Check if a file extension is supported."""
    return extension.lower() in FileExtensions.__members__.values()